"""
HUMANIZADOR DE TEXTO
"""

#!/usr/bin/env python
# coding: utf-8

# In[ ]:





# In[6]:


import os
from docx import Document
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain_openai.chat_models import ChatOpenAI
from PyPDF2 import PdfReader
import nltk
from nltk.tokenize import sent_tokenize
from langchain_openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
import concurrent.futures

# Descargar recursos de NLTK (si es la primera vez)
# nltk.download('punkt')

# =============================================================================
# CONFIGURACIÓN Y CLIENTES
# =============================================================================

load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise ValueError("La clave de la API de OpenAI no se encontró en las variables de entorno.")

llm = ChatOpenAI(api_key=API_KEY, model="o3-mini-2025-01-31", temperature=1)
embeddings = OpenAIEmbeddings(api_key=API_KEY)

# =============================================================================
# TEMPLATES DE PROMPT (para generar una única oración)
# =============================================================================

PROMPT_ESTILO = PromptTemplate(
    input_variables=["texto_estilo", "contenido"],
    template="""Ajusta tu forma de redacción para cumplir con los siguientes criterios, asegurando un texto que se asemeje a una respuesta completamente humana y alejada de los patrones de escritura típicos de inteligencia artificial.

Estilo de Redacción:
Imita la fluidez, la coherencia y la sofisticación de un texto académico extraído de un documento formal en español. Para ello, ten en cuenta las siguientes directrices:

Extensión y Complejidad de las Frases:
Redacta oraciones extensas que integren varias ideas de manera armónica, utilizando conectores fluidos en español.

Estructura Cohesiva y Fragmentación Adecuada:
Organiza el contenido en un mínimo de dos y un máximo de tres oraciones interconectadas, asegurando una redacción natural y evitando frases inconexas.

Claridad y Consistencia:
Cada oración debe mantener un balance entre precisión y fluidez, evitando repeticiones innecesarias o explicaciones redundantes.

Uso Estratégico de Conectores:
Minimiza los puntos aislados y favorece una conexión fluida de ideas mediante conjunciones variadas, sin caer en estructuras mecánicas.

Diversificación Sintáctica:
Utiliza una combinación de oraciones subordinadas, construcciones con participios y elementos discursivos que den variedad y riqueza al texto.

Tono Académico pero Accesible:
Mantén un estilo formal y estructurado, pero sin tecnicismos excesivos que dificulten la comprensión. Piensa en la redacción de un libro universitario.

Instrucciones Finales:
Redacta el siguiente contenido aplicando todas las reglas anteriores, asegurando que el resultado sea indistinguible de un texto humano de alto nivel académico. Evita frases cortas, patrones predecibles o un tono excesivamente robótico. Usa la referencia proporcionada a continuación como base para emular el estilo lingüístico y discursivo.

Ejemplos de redacción: Utiliza de ejemplo las siguientes frases {texto_estilo}

Referencia de Estilo:
{texto_estilo}

Texto de Entrada:
{contenido}

Salida esperada:
Una única oración que preserve el sentido original y se presente con un estilo humanizado y académico."""
)

# =============================================================================
# FUNCIONES DE UTILIDAD
# =============================================================================

def get_pdf_text(ruta_pdf: str, max_tokens: int = 100000) -> str:
    """
    Extrae el texto de un PDF, limitándolo a max_tokens tokens.
    """
    try:
        reader = PdfReader(ruta_pdf)
        tokens_totales = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                tokens_totales.extend(page_text.split())
                if len(tokens_totales) >= max_tokens:
                    tokens_totales = tokens_totales[:max_tokens]
                    break
        print(f"Total tokens leídos del PDF: {len(tokens_totales)}")
        return " ".join(tokens_totales)
    except Exception as e:
        print(f"Error al leer el PDF: {e}")
        return ""

def chunk_text_to_sentences(texto: str, min_words: int = 5) -> list:
    """
    Separa el texto en oraciones y retorna aquellas que tengan al menos min_words.
    """
    oraciones = sent_tokenize(texto, language='spanish')
    return [oracion.strip() for oracion in oraciones if len(oracion.split()) >= min_words]

def create_style_vectorstore(ruta_pdf_estilo: str) -> FAISS:
    """
    Crea una base vectorial FAISS a partir del PDF de estilo, segmentándolo en oraciones.
    """
    texto_pdf = get_pdf_text(ruta_pdf_estilo)
    oraciones = chunk_text_to_sentences(texto_pdf, min_words=10)
    if not oraciones:
        raise ValueError("No se encontraron oraciones válidas en el PDF de estilo.")
    return FAISS.from_texts(oraciones, embeddings)

def retrieve_style_text(vectorstore: FAISS, query: str = "Extrae las oraciones que mejor representen un estilo formal, académico y fluido", k: int = 10) -> str:
    """
    Recupera de la base vectorial las oraciones que mejor representan el estilo deseado.
    """
    resultados = vectorstore.similarity_search(query, k=k)
    oraciones = [doc.page_content for doc in resultados]
    return "\n".join(oraciones)

def apply_style_to_text(texto: str, texto_estilo: str) -> str:
    """
    Aplica la humanización al texto usando el estilo definido en 'texto_estilo'
    y genera una única oración que conserve el sentido original.
    """
    chain = PROMPT_ESTILO | llm
    response = chain.invoke({
        "texto_estilo": texto_estilo,
        "contenido": texto
    })
    return response.content.strip()

def process_paragraph(paragraph_text: str, texto_estilo: str) -> str:
    """
    Divide el párrafo en oraciones y, para cada una:
      - Si la oración tiene 15 palabras o menos, se copia sin modificar.
      - De lo contrario, se humaniza de forma concurrente.
    Luego, se reensambla el párrafo en una única cadena manteniendo el orden original.
    """
    sentences = sent_tokenize(paragraph_text, language='spanish')
    if not sentences:
        return ""

    # Inicializar la lista de resultados con el mismo tamaño que las oraciones
    results = [None] * len(sentences)
    futures = {}

    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        for idx, sentence in enumerate(sentences):
            if len(sentence.split()) <= 15:
                results[idx] = sentence  # Se copia sin modificar
            else:
                future = executor.submit(apply_style_to_text, sentence, texto_estilo)
                futures[future] = idx
        for future in concurrent.futures.as_completed(futures):
            idx = futures[future]
            results[idx] = future.result()
    return " ".join(results)

# =============================================================================
# FUNCIÓN PARA GENERAR UN DOCUMENTO HUMANIZADO MANTENIENDO LA ESTRUCTURA
# =============================================================================

def generate_humanized_word_document(input_path: str, output_path: str, ruta_pdf_estilo: str):
    """
    Lee un documento Word de entrada, preserva la estructura (títulos, subtítulos, secciones)
    y procesa únicamente los párrafos "normales" (aquellos que no sean Titulo1_Cato o Titulo2_Cato)
    aplicando la humanización a cada oración individualmente (salvo las oraciones de 15 palabras o menos),
    reconociendo los estilos:
      - "Titulo1_Cato" (equivalente a título)
      - "Titulo2_Cato" (equivalente a subtítulo)
    Los demás párrafos se procesan o se copian según su estilo original.
    """
    # Crear vectorstore y obtener el texto de estilo
    vectorstore = create_style_vectorstore(ruta_pdf_estilo)
    texto_estilo = retrieve_style_text(vectorstore)

    # Abrir el documento de entrada
    input_doc = Document(input_path)
    output_doc = Document()

    # Iterar por cada párrafo y copiar su contenido y estilo
    for p in input_doc.paragraphs:
        if not p.text.strip():
            continue

        style_name = p.style.name
        # Si el párrafo es Titulo1_Cato o Titulo2_Cato o un encabezado reconocido, se copia sin modificación.
        if style_name in ["Titulo1_Cato", "Heading 1"]:
            output_doc.add_paragraph(p.text, style="Heading 1")
        elif style_name in ["TITULO2_Cato", "Heading 2"]:
            output_doc.add_paragraph(p.text, style="Heading 2")
        elif style_name.startswith("Heading"):
            output_doc.add_paragraph(p.text, style=style_name)
        else:
            # Para párrafos "normales" se procesa el texto.
            processed_text = process_paragraph(p.text, texto_estilo)
            output_doc.add_paragraph(processed_text, style=style_name)

    output_doc.save(output_path)
    print(f"Documento humanizado guardado en: {output_path}")

# =============================================================================
# BLOQUE PRINCIPAL
# =============================================================================

if __name__ == "__main__":
    input_word_path = r"C:\Users\HP\Desktop\LIBROS PERSO\PROPUESTA DE INVESTIGACION\PROYECTO DE INVESTIGACIÓN.docx"
    output_doc_path = r"C:\Users\HP\Desktop\LIBROS PERSO\PROPUESTA DE INVESTIGACION\PROYECTO_DE_INVESTIGACIÓN_humanizado.docx"
    ruta_pdf_estilo = r"C:\Users\HP\Desktop\LIBROS PERSO\CONTEXTO ESPANIOL\CONTEXTO5.pdf"

    generate_humanized_word_document(input_word_path, output_doc_path, ruta_pdf_estilo)


# In[ ]:




